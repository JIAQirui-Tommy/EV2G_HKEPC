from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("/Users/tommy/Desktop/EV2Gym-main")
OUT = ROOT / "HKEPC_2026_reformatted_paper_condensed.docx"


def set_run_font(run, *, size=10, bold=False, italic=False, superscript=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.superscript = superscript


def add_paragraph(doc, text="", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, before=0, after=2):
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.line_spacing = 1.0
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if indent:
        fmt.first_line_indent = Cm(0.63)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_heading(doc, text, *, before=6, after=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.line_spacing = 1.0
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, bold=True)
    return p


def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.line_spacing = 1.0
    fmt.space_before = Pt(1)
    fmt.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.line_spacing = 1.0
    fmt.space_before = Pt(1)
    fmt.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_image(doc, path, caption, width_cm=12.6):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.line_spacing = 1.0
    fmt.space_before = Pt(2)
    fmt.space_after = Pt(1)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def set_cell(cell, text, *, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.line_spacing = 1.0
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, bold=bold)


def add_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "auto")
        borders.append(elem)
    tbl_pr.append(borders)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
section.header_distance = Cm(1.25)
section.footer_distance = Cm(1.25)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
style.font.size = Pt(10)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.0
p.paragraph_format.space_after = Pt(4)
run = p.add_run(
    "Reward Design Matters: An AI-Based Wireless EV\n"
    "Charging Prototype for User Satisfaction and Peak\n"
    "Shaving in Smart Grids"
)
set_run_font(run, size=12, bold=True)

# Authors
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.0
p.paragraph_format.space_after = Pt(1)
for i, name in enumerate(["ZHANG Tianyi", "JIA Qirui", "KUANG JINLI"]):
    if i:
        sep = p.add_run(", ")
        set_run_font(sep)
    r = p.add_run(name)
    set_run_font(r)
    s = p.add_run("1")
    set_run_font(s, superscript=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.0
p.paragraph_format.space_after = Pt(4)
r = p.add_run("1 The Hong Kong Polytechnic University, Hong Kong, China, Email: [please add email]")
set_run_font(r)

add_heading(doc, "ABSTRACT", before=0)
add_paragraph(
    doc,
    "This paper presents an AI-based wireless electric vehicle (EV) charging prototype for smart grid and smart city applications. The framework integrates wireless charging zones, dynamic electricity prices, user state-of-charge (SoC) requirements, transformer constraints, and a forecasting-assisted control pipeline. A PPO-based reinforcement learning (RL) controller is compared with a heuristic baseline. Results show that RL maintains feasible operation and high user satisfaction, but remains economically conservative, with limited discharge and weak price-following behavior. The experiments indicate that reward design is the key bottleneck: a satisfaction-dominant reward promotes safe operation but discourages profitable discharge. The framework is therefore positioned as a practical early-stage prototype and a step toward smarter wireless EV charging in future urban energy systems.",
    indent=False,
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.line_spacing = 1.0
p.paragraph_format.space_after = Pt(3)
r = p.add_run("Index Terms")
set_run_font(r, bold=True)
r = p.add_run("—wireless EV charging, reinforcement learning, smart grid, smart city, reward design, peak shaving")
set_run_font(r)

add_heading(doc, "I. INTRODUCTION")
for text in [
    "Electric vehicle charging is becoming an important part of future smart cities. A charging system must satisfy user demand, protect grid assets, and respond to electricity prices. These goals often conflict: a controller focused only on completion may ignore economic value, while one focused only on profit may fail to deliver the target battery level before departure.",
    "Wireless charging adds practical significance because charging and discharging can occur with less user intervention. This supports a smoother energy-management process and may improve user acceptance in future smart city deployments.",
    "This paper studies that problem using a simulation-driven AI framework with dynamic prices and transformer constraints. The main conclusion is not simply that RL works, but that the current RL controller learns safe behavior while remaining economically conservative because of reward bias.",
    "The contributions are: 1) a wireless EV charging architecture with six wireless charging zones, each serving eight EV slots, under transformer-aware control and dynamic pricing; 2) identification of a key RL failure mode, where a satisfaction- and safety-dominant reward suppresses economically useful discharge; 3) a forecasting-assisted evaluation pipeline combining price response, peak shaving, and user completion analysis; and 4) a practical improvement direction based on dynamic reward design.",
]:
    add_paragraph(doc, text)

add_heading(doc, "II. SYSTEM OVERVIEW AND PROBLEM FORMULATION")
for text in [
    "The framework has four layers: data, forecasting, simulation, and control. The data layer includes price signals, EV battery parameters, target SoC, and transformer limits. The forecasting layer provides short-horizon predictive information. The simulation layer models EV arrival, departure, charging, discharging, and transformer loading. The control layer compares RL with a heuristic baseline. The prototype extends the open-source EV2Gym platform [1] with wireless charging zones, forecasting inputs, and a Hong Kong-oriented proxy scenario.",
    "The wireless charging setting contains six wireless charging zones, each with eight EV positions, so the controller manages up to forty-eight EV charging interfaces. Wireless charging matters because it enables charging and discharging with minimal user action, supporting smoother background energy management and more realistic peak shaving in future smart city systems.",
    "Direct public operational data for Hong Kong remain limited. Therefore, the current prototype uses a proxy scenario that combines local charging assumptions with Greater Bay Area electricity-use preference patterns and regional price information to approximate a Hong Kong-like environment.",
    "At each time step t, the controller observes a state s_t and selects an action a_t. The state includes time, EV SoC, remaining departure time, price information, and aggregate power. The action is the charging or discharging power assigned to each wireless zone. The objective is to maximize long-term return while maintaining user service and grid feasibility:",
]:
    add_paragraph(doc, text)

add_equation(doc, "max π  E [ Σ(T-1,t=0) γ^t r_t ]    (1)")
add_paragraph(
    doc,
    "Here, π is the control policy, γ is the discount factor, and r_t is the reward at time step t. This means the controller is trained to maximize total discounted reward over the charging horizon."
)

add_heading(doc, "III. METHODOLOGY")
add_heading(doc, "A. RL-Based Charging Control", before=2)
for text in [
    "The RL controller uses PPO. The state includes current SoC, target SoC, departure urgency, current power, and price information. In the wireless setting, a forecast signal is also included to provide limited look-ahead support. The action determines charging or discharging intensity in each wireless zone.",
]:
    add_paragraph(doc, text)

add_equation(doc, "R_t = αS_t + βP_t + γG_t    (2)")
for text in [
    "In Eq. (2), S_t is the satisfaction term, P_t is the profit term, and G_t is the grid-safety term. The parameter α controls user-completion emphasis, β controls economic return under dynamic prices, and γ controls avoidance of unsafe grid behavior.",
    "This formula explains the current failure mode. In the present prototype, α and the safety-related part are effectively too strong compared with β. As a result, the RL agent learns that preserving SoC is safer than discharging for price gain. The issue is not lack of price visibility, but an objective that values missed user demand as much more costly than missed economic opportunity.",
]:
    add_paragraph(doc, text)

add_heading(doc, "B. Forecasting Module", before=2)
for text in [
    "A Prophet-style forecasting module provides a short-horizon predictive layer. Its role is not to replace the controller, but to summarize future load movement so that the RL policy does not rely only on the current time step. This is important when the controller must decide whether to keep charging, slow down, or reserve energy for later peak shaving and valley filling.",
]:
    add_paragraph(doc, text)

add_equation(doc, "y_t = g_t + s_t + h_t + ε_t    (3)")
add_paragraph(
    doc,
    "Here, g_t is the long-term trend, s_t is the seasonal component, h_t is the holiday or event effect, and ε_t is the residual error. In this project, these terms describe gradual urban load change, repeated charging patterns, predictable abnormal demand, and short-term fluctuation, respectively."
)
add_equation(doc, "s_t = Σ(N,n=1) [ a_n cos(2πnt/P) + b_n sin(2πnt/P) ]    (4)")
add_paragraph(
    doc,
    "In Eq. (4), a_n and b_n are Fourier coefficients, P is the period, and N is the order. The forecasting block is used at both daily and hourly levels so that the controller can track broad demand trends and local turning points under dynamic pricing."
)

add_image(
    doc,
    ROOT / "results/wireless_rl_sim_2026_03_28_550433/Transformer_Aggregated_Power.png",
    "Fig. 1: Transformer-level power allocation under the 100 kW limit.",
    width_cm=11.5,
)

add_heading(doc, "C. Dynamic Reward Design", before=2)
add_equation(doc, "R_t = α_tS_t + β_tP_t + γG_t    (5)")
add_paragraph(
    doc,
    "To address the conservative behavior, α_t and β_t are allowed to change over time. When the price spread is high, β_t should increase so that the controller values profit more. When departure time is near, α_t should increase so that user satisfaction remains dominant. This is a simple and practical improvement because it changes reward balance without requiring a new learning algorithm."
)

add_heading(doc, "IV. EXPERIMENTAL RESULTS")
add_heading(doc, "A. Wireless Charging Operation and Peak Shaving", before=2)
for text in [
    "The wireless charging results show why this setting matters. The transformer-side power profile stays below the 100 kW limit, indicating that the controller respects transformer constraints. More importantly, the total profile changes sign across the day, showing that the system is capable of both charging and discharging, which is the basis of peak shaving and valley filling.",
    "Zone-level current and SoC profiles also show that vehicles are served across multiple zones over time. This supports the practical value of wireless charging because energy management can happen in the background without requiring users to manually intervene at each step.",
]:
    add_paragraph(doc, text)

add_heading(doc, "B. Price Response and RL-vs-Heuristic Comparison", before=2)
for text in [
    "The charge and discharge price curves contain clear low-price and high-price periods. A good controller should charge more during low-price intervals and discharge more when prices rise. The heuristic follows this structure more clearly, while the RL controller remains conservative.",
    "The metric comparison confirms this observation. RL keeps satisfaction high and produces substantial discharged energy, but the heuristic still achieves better profit in this scenario. The comparison suggests that RL has not yet converted forecast and price information into the strongest economic response.",
]:
    add_paragraph(doc, text)

add_image(
    doc,
    ROOT / "results/wireless_compare_2026_03_28_221326/wireless_power_price_compare.png",
    "Fig. 2: RL versus heuristic power behavior with dynamic price signals.",
    width_cm=11.5,
)

add_heading(doc, "C. Improvement Experiment: Dynamic Reward Perspective", before=2)
add_paragraph(
    doc,
    "To address the reward bottleneck, a design-oriented improvement experiment based on dynamic reward weighting is proposed. The original RL gives high satisfaction but low profit. The heuristic gives a medium-level balance. The improved RL should preserve high satisfaction while increasing profit response."
)

table = doc.add_table(rows=4, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_borders(table)
for idx, val in enumerate(["Method", "Profit", "Satisfaction"]):
    set_cell(table.cell(0, idx), val, bold=True)
for r, row in enumerate(
    [
        ("Original RL", "Low", "High"),
        ("Heuristic", "Medium", "Medium"),
        ("Improved RL", "High", "High"),
    ],
    start=1,
):
    for c, val in enumerate(row):
        set_cell(table.cell(r, c), val)
add_caption(doc, "Table I: Target comparison for the improved controller.")

add_heading(doc, "V. DISCUSSION")
for text in [
    "The main value of this work is that it identifies a clear mechanism. The RL controller is not random or useless: it learns stable operation, keeps user satisfaction high, respects the transformer limit, and shows some charge-discharge ability. These are meaningful results for a wireless charging prototype.",
    "However, the current RL policy does not yet deliver the strongest peak-shaving and valley-filling behavior. The power curves show that the system can shift load, but the economic timing remains weak. The reason is reward bias. The current reward values user completion more than profitable discharge, so the policy keeps extra energy margin instead of using it when price signals suggest it should.",
    "This finding is important because it is specific and actionable. It suggests that the present limitation is not simply lack of training or poor algorithm choice, but a mismatch between the control objective and the desired economic behavior. Limited Hong Kong public data remain a constraint, but reward redesign is still the main step required to unlock stronger economic performance.",
]:
    add_paragraph(doc, text)

add_heading(doc, "VI. COMMERCIAL VALUE AND SMART CITY RELEVANCE")
for text in [
    "This system has direct smart city value. Wireless charging makes energy exchange less visible to users, so charging and discharging can happen more smoothly in the background. That is important for future urban fleets, public charging facilities, and integrated mobility-energy systems.",
    "The commercial value is also clear. A better controller can reduce transformer stress, support peak shaving, improve charging service, and increase station-level operational efficiency. For utilities and city-scale aggregators, this framework can help coordinate many EVs as flexible energy resources rather than treating them only as passive loads.",
]:
    add_paragraph(doc, text)

add_heading(doc, "VII. CONCLUSION AND FUTURE WORK")
for text in [
    "This paper presented an AI-based wireless EV charging prototype with transformer-aware control, Prophet-style forecasting, PPO-based RL, and heuristic comparison. The system is functional and shows that wireless charging can support background charge-discharge scheduling for smart grid applications.",
    "The current RL controller achieves feasible operation and high satisfaction, but still underuses price-responsive discharge and does not fully realize peak-shaving and valley-filling value. The main conclusion is therefore that reward design is the key bottleneck. Future work should improve dynamic reward weighting, strengthen explicit profit signals, and validate the controller with richer real-world data.",
]:
    add_paragraph(doc, text)

add_heading(doc, "REFERENCES")
refs = [
    "[1] S. Orfanos, “EV2Gym,” GitHub repository. [Online]. Available: https://github.com/StavrosOrf/EV2Gym. Accessed: Mar. 30, 2026.",
    "[2] Y. Li, H. Zhang, X. Chen, and Y. Wang, “Reinforcement learning for electric vehicle charging scheduling: A systematic review,” Transportation Research Part E: Logistics and Transportation Review, vol. 190, p. 103698, 2024.",
    "[3] M. K. M. van der Meer, J. G. Slootweg, and M. Gibescu, “Development and evaluation of a smart charging strategy for an electric vehicle fleet based on reinforcement learning,” Applied Energy, vol. 285, p. 116382, 2021.",
    "[4] R. S. Sutton and A. G. Barto, Reinforcement Learning: An Introduction, 2nd ed. Cambridge, MA, USA: MIT Press, 2018.",
]
for ref in refs:
    add_paragraph(doc, ref, indent=False, after=1)

doc.save(OUT)
print(f"Wrote {OUT}")
